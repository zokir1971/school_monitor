from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document
from sqlalchemy import select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.ext.asyncio import AsyncSession

from app.modules.planning.models import PlanTemplate, PlanDirection, PlanTemplateRow4, PlanTemplateRow11, \
    PlanTemplateDirection
from app.modules.planning.repo import (
    PlanDirectionRepo,
    PlanTemplateRepo,
    PlanTemplateRow4Repo,
    PlanTemplateRow11Repo,
    PlanTemplateDirectionRepo,
    PlanTemplateReplaceRepo
)


# ---------------------------------------------------------
# Виртуальный "section", чтобы твой Jinja-шаблон работал
# (он ожидает sections и section.id)
# ---------------------------------------------------------
@dataclass(frozen=True)
class VirtualSection:
    id: int  # будем использовать direction_id
    title: str  # для отображения (если нужно)


@dataclass(frozen=True)
class TemplatesPageData:
    selected_direction: Any | None
    selected_direction_id: int | None
    templates: list[Any]
    selected_template_id: int | None
    sections: list[VirtualSection]
    items_by_section: dict[int, list[Any]]


class PlanningSuperService:
    """
    Сервис для страницы "Шаблон ВШК" (superuser).
    Реальные строки берём из PlanTemplateRow4 (таблица plan_template_rows_4).
    """

    # =========================================================
    # Page builder (данные для рендера)
    # =========================================================
    @staticmethod
    async def build_page(
            db: AsyncSession,
            *,
            direction_id: int | None,
            template_id: int | None,
    ) -> TemplatesPageData:
        # 1) направление
        selected_direction = None
        if direction_id:
            selected_direction = await PlanDirectionRepo.get(db, direction_id)

        # 2) шаблоны (глобальные, не привязаны к направлению)
        templates = await PlanTemplateRepo.list_all(db)

        # 3) выбранный шаблон:
        selected_template_id = template_id
        if selected_template_id is None and templates:
            # если не передали template_id, возьмём активный или последний
            active_or_latest = await PlanTemplateRepo.get_active_or_latest(db)
            selected_template_id = active_or_latest.id if active_or_latest else templates[0].id

        # 4) sections: чтобы шаблон работал — создаём 1 виртуальный section под направление
        sections: list[VirtualSection] = []
        items_by_section: dict[int, list[Any]] = {}

        if selected_direction and selected_template_id:
            sec_title = selected_direction.full_title or selected_direction.short_title
            sections = [VirtualSection(id=selected_direction.id, title=sec_title)]

            rows = await PlanTemplateRow4Repo.list_rows(
                db,
                template_id=selected_template_id,
                direction_id=selected_direction.id,
            )
            items_by_section[selected_direction.id] = list(rows)

        return TemplatesPageData(
            selected_direction=selected_direction,
            selected_direction_id=direction_id,
            templates=list(templates),
            selected_template_id=selected_template_id,
            sections=sections,
            items_by_section=items_by_section,
        )

    # =========================================================
    # Actions
    # =========================================================
    @staticmethod
    async def create_template(
            db: AsyncSession,
            *,
            name: str,
            academic_year: str | None,
            created_by_user_id: int | None,
    ) -> int:
        """
        Создать PlanTemplate.
        """
        tpl = await PlanTemplateRepo.create(
            db,
            name=name,
            academic_year=academic_year,
            created_by_user_id=created_by_user_id,
            is_active=False,
        )
        return tpl.id

    @staticmethod
    async def add_row4(
            db: AsyncSession,
            *,
            template_id: int,
            direction_id: int,
            control_object: str,
            risks: str | None,
            decisions: str | None,
            row_order: int | None,
    ) -> int:
        control_object_val = (control_object or "").strip()
        if not control_object_val:
            raise ValueError("Поле 'Бақылау нысаны' обязательно.")

        # 1) если не задано — добавляем в конец
        if row_order is None or int(row_order) <= 0:
            mx = await PlanTemplateRow4Repo.max_row_order(
                db, template_id=template_id, direction_id=direction_id
            )
            row_order_val = mx + 1
        else:
            row_order_val = int(row_order)

        # 2) пытаемся вставить с указанным row_order
        try:
            row = await PlanTemplateRow4Repo.create(
                db,
                template_id=template_id,
                direction_id=direction_id,
                row_order=row_order_val,
                control_object=control_object_val,
                risk_text=(risks or "").strip() or None,
                decision_text=(decisions or "").strip() or None,
                no=None,
            )
            return row.id

        except IntegrityError:
            # если номер занят — добавим в конец
            await db.rollback()  # важно: очистить failed transaction

            mx = await PlanTemplateRow4Repo.max_row_order(
                db, template_id=template_id, direction_id=direction_id
            )
            row = await PlanTemplateRow4Repo.create(
                db,
                template_id=template_id,
                direction_id=direction_id,
                row_order=mx + 1,
                control_object=control_object_val,
                risk_text=(risks or "").strip() or None,
                decision_text=(decisions or "").strip() or None,
                no=None,
            )
            return row.id

    @staticmethod
    async def delete_row4(db: AsyncSession, *, row_id: int) -> None:
        ok = await PlanTemplateRow4Repo.delete_row(db, row_id=row_id)
        if not ok:
            raise ValueError("Строка не найдена или уже удалена.")

    # =========================================================
    # Import .docx (REPLACE) в rows_4
    # =========================================================
    @staticmethod
    async def import_docx_row4_replace(
            db: AsyncSession,
            *,
            template_id: int,
            direction_id: int,
            file_bytes: bytes,
    ) -> int:
        """
        Импортирует первую таблицу .docx в plan_template_rows_4 (REPLACE)
        для пары (template_id, direction_id).

        Первая строка таблицы считается заголовком и НЕ сохраняется.
        """

        # 1) удаляем старые строки
        await PlanTemplateRow4Repo.delete_all_for_template_direction(
            db,
            template_id=template_id,
            direction_id=direction_id,
        )

        # 2) парсим docx
        doc = Document(BytesIO(file_bytes))
        if not doc.tables:
            raise ValueError("В документе нет таблиц (.docx).")

        table = doc.tables[0]
        if not table.rows:
            return 0

        imported = 0

        for r_idx, row in enumerate(table.rows):
            if r_idx == 0:
                continue  # заголовок

            cells = row.cells
            if len(cells) < 3:
                continue

            # --- Вариант А: docx с 4 колонками (№ + 3 поля)
            if len(cells) >= 4:
                no = (cells[0].text or "").strip() or None
                control_object = (cells[1].text or "").strip()
                risk_text = (cells[2].text or "").strip() or None
                decision_text = (cells[3].text or "").strip() or None
            # --- Вариант B: docx с 3 колонками (без №)
            else:
                no = None
                control_object = (cells[0].text or "").strip()
                risk_text = (cells[1].text or "").strip() or None
                decision_text = (cells[2].text or "").strip() or None

            if not control_object:
                continue

            await PlanTemplateRow4Repo.create(
                db,
                template_id=template_id,
                direction_id=direction_id,
                row_order=imported + 1,
                no=no,  # ✅ сохраним номер из Word (если есть)
                control_object=control_object,
                risk_text=risk_text,
                decision_text=decision_text,
            )
            imported += 1


@dataclass(frozen=True)
class VirtualSection:
    id: int
    title: str


@dataclass(frozen=True)
class Templates11PageData:
    selected_direction: Any | None
    selected_direction_id: int | None
    templates: list[Any]
    selected_template_id: int | None
    sections: list[VirtualSection]
    items_by_section: dict[int, list[Any]]


class PlanningSuperTemplates11Service:
    @staticmethod
    async def build_page(
            db: AsyncSession,
            *,
            direction_id: int | None,
            template_id: int | None,
    ) -> Templates11PageData:
        selected_direction = await PlanDirectionRepo.get(db, direction_id) if direction_id else None
        templates = await PlanTemplateRepo.list_all(db)

        selected_template_id = template_id
        if selected_template_id is None and templates:
            active_or_latest = await PlanTemplateRepo.get_active_or_latest(db)
            selected_template_id = active_or_latest.id if active_or_latest else templates[0].id

        sections: list[VirtualSection] = []
        items_by_section: dict[int, list[Any]] = {}

        if selected_direction and selected_template_id:
            title = selected_direction.full_title or selected_direction.short_title
            sections = [VirtualSection(id=selected_direction.id, title=title)]

            rows = await PlanTemplateRow11Repo.list_rows(
                db,
                template_id=selected_template_id,
                direction_id=selected_direction.id,
            )
            items_by_section[selected_direction.id] = list(rows)

        return Templates11PageData(
            selected_direction=selected_direction,
            selected_direction_id=direction_id,
            templates=list(templates),
            selected_template_id=selected_template_id,
            sections=sections,
            items_by_section=items_by_section,
        )

    @staticmethod
    async def add_row11(
            db: AsyncSession,
            *,
            template_id: int,
            direction_id: int,
            topic: str,
            goal: str | None,
            control_object: str | None,
            control_type: str | None,
            methods: str | None,
            deadlines: str | None,
            responsibles: str | None,
            review_place: str | None,
            management_decision: str | None,
            second_control: str | None,
            row_order: int | None,
    ) -> int:
        topic_val = (topic or "").strip()
        if not topic_val:
            raise ValueError("Поле 'Бақылау тақырыбы' обязательно.")

        # безопасный row_order (из-за UniqueConstraint template_id+direction_id+row_order)
        if row_order is None or int(row_order) <= 0:
            mx = await PlanTemplateRow11Repo.max_row_order(db, template_id=template_id, direction_id=direction_id)
            row_order_val = mx + 1
        else:
            row_order_val = int(row_order)

        try:
            row = await PlanTemplateRow11Repo.create(
                db,
                template_id=template_id,
                direction_id=direction_id,
                row_order=row_order_val,
                topic=topic_val,
                goal=goal,
                control_object=control_object,
                control_type=control_type,
                methods=methods,
                deadlines=deadlines,
                responsibles=responsibles,
                review_place=review_place,
                management_decision=management_decision,
                second_control=second_control,
            )
            return row.id
        except IntegrityError:
            await db.rollback()
            mx = await PlanTemplateRow11Repo.max_row_order(db, template_id=template_id, direction_id=direction_id)
            row = await PlanTemplateRow11Repo.create(
                db,
                template_id=template_id,
                direction_id=direction_id,
                row_order=mx + 1,
                topic=topic_val,
                goal=goal,
                control_object=control_object,
                control_type=control_type,
                methods=methods,
                deadlines=deadlines,
                responsibles=responsibles,
                review_place=review_place,
                management_decision=management_decision,
                second_control=second_control,
            )
            return row.id

    @staticmethod
    async def delete_row11(db: AsyncSession, *, row_id: int) -> None:
        ok = await PlanTemplateRow11Repo.delete_row(db, row_id=row_id)
        if not ok:
            raise ValueError("Строка не найдена или уже удалена.")

    @staticmethod
    async def import_docx_row11_replace(
            db: AsyncSession,
            *,
            template_id: int,
            direction_id: int,
            file_bytes: bytes,
    ) -> int:
        """
        DOCX: 11 колонок = № + 10 полей.
        Первая строка — заголовок, НЕ сохраняем.
        """
        await PlanTemplateRow11Repo.delete_all_for_template_direction(
            db, template_id=template_id, direction_id=direction_id
        )

        doc = Document(BytesIO(file_bytes))
        if not doc.tables:
            raise ValueError("В документе нет таблиц (.docx).")

        table = doc.tables[0]
        if not table.rows:
            return 0

        imported = 0

        for r_idx, row in enumerate(table.rows):
            if r_idx == 0:
                continue  # заголовок

            cells = row.cells
            if len(cells) < 11:
                # ожидаем № + 10 = 11
                continue

            topic = (cells[1].text or "").strip()
            goal = (cells[2].text or "").strip() or None
            control_object = (cells[3].text or "").strip() or None
            control_type = (cells[4].text or "").strip() or None
            methods = (cells[5].text or "").strip() or None
            deadlines = (cells[6].text or "").strip() or None
            responsibles = (cells[7].text or "").strip() or None
            review_place = (cells[8].text or "").strip() or None
            management_decision = (cells[9].text or "").strip() or None
            second_control = (cells[10].text or "").strip() or None

            if not topic:
                continue

            await PlanTemplateRow11Repo.create(
                db,
                template_id=template_id,
                direction_id=direction_id,
                row_order=imported + 1,
                topic=topic,
                goal=goal,
                control_object=control_object,
                control_type=control_type,
                methods=methods,
                deadlines=deadlines,
                responsibles=responsibles,
                review_place=review_place,
                management_decision=management_decision,
                second_control=second_control,
            )
            imported += 1

        return imported


def _cell_text(cell):
    # объединяем параграфы через перенос строки
    parts = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts).strip()


def _table_to_matrix(table):
    matrix = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(_cell_text(cell))
        matrix.append(row_data)
    return matrix


def _drop_header_if_needed(rows: list[list[str]]) -> list[list[str]]:
    # Если первая строка похожа на заголовок — убираем её.
    if not rows:
        return rows
    first = " ".join(rows[0]).lower()
    header_markers = ["№", "no", "объект", "тема", "goal", "цель", "responsible", "ответ"]
    if any(m.lower() in first for m in header_markers):
        return rows[1:]
    return rows


def _normalize_none(s: str) -> str | None:
    s = (s or "").strip()
    return s or None


def _fix_docx_shift_for_row11(row: list[str]) -> list[str]:
    # убираем пустые хвосты
    row = [(c or "").strip() for c in row]

    # Ситуация: ["1", "1", "Бақылау тақырыбы...", ...]
    # т.е. номер продублировался -> удаляем один из них
    if len(row) >= 2 and row[0].isdigit() and row[1] == row[0]:
        row = row[1:]  # убрали дубль

    return row


class SuperuserBuilderService:
    @staticmethod
    async def build_page(
            db: AsyncSession,
            *,
            template_id: int | None,
            direction_id: int | None,
    ) -> dict:
        templates = (await db.execute(
            select(PlanTemplate).order_by(PlanTemplate.is_active.desc(), PlanTemplate.created_at.desc())
        )).scalars().all()

        directions = (await db.execute(
            select(PlanDirection).order_by(PlanDirection.sort_order.asc(), PlanDirection.id.asc())
        )).scalars().all()

        rows4: list[PlanTemplateRow4] = []
        rows11: list[PlanTemplateRow11] = []

        if template_id and direction_id:
            td = await db.scalar(
                select(PlanTemplateDirection).where(
                    PlanTemplateDirection.template_id == template_id,
                    PlanTemplateDirection.direction_id == direction_id,
                )
            )
            if td:
                res4 = await db.execute(
                    select(PlanTemplateRow4)
                    .where(PlanTemplateRow4.template_direction_id == td.id)
                    .order_by(PlanTemplateRow4.row_order.asc(), PlanTemplateRow4.id.asc())
                )
                rows4: list[PlanTemplateRow4] = list(res4.scalars().all())

                res11 = await db.execute(
                    select(PlanTemplateRow11)
                    .where(PlanTemplateRow11.template_direction_id == td.id)
                    .order_by(PlanTemplateRow11.row_order.asc(), PlanTemplateRow11.id.asc())
                )
                rows11: list[PlanTemplateRow11] = list(res11.scalars().all())
            else:
                rows4: list[PlanTemplateRow4] = []
                rows11: list[PlanTemplateRow11] = []

        return {
            "templates": templates,
            "directions": directions,
            "selected_template_id": template_id,
            "selected_direction_id": direction_id,
            "rows4": rows4,
            "rows11": rows11,
        }

    @staticmethod
    async def import_docx_two_tables_replace(
            db: AsyncSession,
            *,
            template_id: int,
            direction_id: int,
            file_bytes: bytes,
    ) -> tuple[int, int]:
        doc = Document(BytesIO(file_bytes))
        if len(doc.tables) < 2:
            raise ValueError("В файле .docx должно быть минимум 2 таблицы: 1) объекты, 2) задачи.")

        t1 = _table_to_matrix(doc.tables[0])
        t2 = _table_to_matrix(doc.tables[1])

        t1 = _drop_header_if_needed(t1)
        t2 = _drop_header_if_needed(t2)

        # ------------- таблица 1 -> Row4 -------------
        # ожидаем 4 колонки в docx: № | объект контроля | риск | управленческое решение
        # ✅ но "№" мы игнорируем, чтобы не было двойной нумерации
        rows4_data: list[dict] = []
        for i, row in enumerate(t1, start=1):
            cols = (row + ["", "", "", ""])[:4]

            rows4_data.append({
                "row_order": i,
                "control_object": (cols[1] or "").strip(),
                "risk_text": _normalize_none(cols[2]),
                "decision_text": _normalize_none(cols[3]),
            })

        # ------------- таблица 2 -> Row11 -------------
        # ожидаем 10 колонок:
        rows11_data: list[dict] = []
        for i, row in enumerate(t2, start=1):
            row = _fix_docx_shift_for_row11(row)

            # № -> ожидаем 11 колонок и игнорируем первую
            cols = (row + [""] * 11)[:11]

            rows11_data.append({
                "row_order": i,
                "topic": (cols[1] or "").strip(),  # после №
                "goal": _normalize_none(cols[2]),
                "control_object": _normalize_none(cols[3]),
                "control_type": _normalize_none(cols[4]),
                "methods": _normalize_none(cols[5]),
                "deadlines": _normalize_none(cols[6]),
                "responsibles": _normalize_none(cols[7]),
                "review_place": _normalize_none(cols[8]),
                "management_decision": _normalize_none(cols[9]),
                "second_control": _normalize_none(cols[10]),
            })

        td = await PlanTemplateDirectionRepo.get_or_create(
            db,
            template_id=template_id,
            direction_id=direction_id,
        )

        c4 = await PlanTemplateReplaceRepo.replace_rows4(
            db, template_direction_id=td.id, rows=rows4_data
        )
        c11 = await PlanTemplateReplaceRepo.replace_rows11(
            db, template_direction_id=td.id, rows=rows11_data
        )

        return c4, c11
